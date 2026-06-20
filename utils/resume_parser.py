from docx import Document
import re



def extract_resume_text(file):


    doc = Document(file)


    content=[]


    for p in doc.paragraphs:


        if p.text.strip():

            content.append(

                p.text.strip()

            )



    for table in doc.tables:


        for row in table.rows:


            content.append(

                " ".join(

                    cell.text.strip()

                    for cell in row.cells

                )

            )



    return "\n".join(content)






def clean_lines(text):


    lines=[]


    for line in text.split("\n"):


        line=line.strip()


        if line:


            lines.append(line)



    return lines






def extract_between(text,start,end_list):


    pattern="("+"|".join(start)+")"


    end="|".join(end_list)



    result=re.search(

        pattern+r"(.*?)(?="+end+"|$)",

        text,

        re.I|re.S

    )


    if result:


        return result.group(2).strip()



    return ""







def parse_resume(text):


    lines=clean_lines(text)



    resume={}





    # -----------------
    # Basic Information
    # -----------------


    resume["name"]=lines[0] if lines else ""



    email=re.findall(

        r'[\w\.-]+@[\w\.-]+',

        text

    )


    resume["email"]=email[0] if email else ""



    phone=re.findall(

        r'(\+?\d[\d\s\-]{8,})',

        text

    )


    resume["phone"]=phone[0] if phone else ""



    resume["linkedin"]=""

    resume["github"]=""

    resume["location"]=""





    # -----------------
    # Job title
    # -----------------


    resume["job_title"]=""



    for line in lines[:15]:


        if any(

            x in line.lower()

            for x in [

                "developer",

                "engineer",

                "analyst",

                "support"

            ]

        ):


            resume["job_title"]=line

            break







    # -----------------
    # Summary
    # -----------------


    summary=extract_between(

        text,

        [

            "PROFESSIONAL SUMMARY",

            "SUMMARY"

        ],

        [

            "CORE SKILLS",

            "TECHNICAL SKILLS",

            "PROFESSIONAL EXPERIENCE",

            "EMPLOYMENT HISTORY"

        ]

    )


    resume["summary"]=summary






    # -----------------
    # Skills
    # -----------------


    skills=extract_between(

        text,

        [

            "CORE SKILLS",

            "TECHNICAL PROFICIENCIES",

            "SKILLS"

        ],

        [

            "PROFESSIONAL EXPERIENCE",

            "EMPLOYMENT HISTORY",

            "PROJECT"

        ]

    )



    resume["skills"]=[


        x.strip()

        for x in re.split(

            ",|\n|\|",

            skills

        )

        if x.strip()

    ]









    # -----------------
    # Experience
    # -----------------


    exp_section=extract_between(

        text,

        [

        "PROFESSIONAL EXPERIENCE",

        "EMPLOYMENT HISTORY"

        ],

        [

        "EDUCATION",

        "ACADEMIC",

        "PROJECT"

        ]

    )



    experiences=[]



    blocks=re.split(

        r'(?=Role\s*:)',

        exp_section,

        flags=re.I

    )





    for block in blocks:


        if len(block)<20:

            continue



        role=re.search(

            r'Role\s*:\s*(.*)',

            block,

            re.I

        )


        client=re.search(

            r'Client\s*:\s*(.*)',

            block,

            re.I

        )


        duration=re.search(

            r'Duration\s*:\s*(.*)',

            block,

            re.I

        )



        responsibilities=[]



        for line in block.split("\n"):


            if len(line)>40 and ":" not in line:


                responsibilities.append(

                    line

                )





        experiences.append(

            {


            "role":

            role.group(1) if role else "",


            "company":

            client.group(1) if client else "",


            "duration":

            duration.group(1) if duration else "",



            "responsibilities":

            responsibilities


            }

        )





    resume["experience"]=experiences








    # -----------------
    # Education
    # -----------------


    education=extract_between(

        text,

        [

        "ACADEMIC DETAILS",

        "EDUCATION"

        ],

        [

        "PROJECT",

        "CERTIFICATION"

        ]

    )



    resume["education"]=[


        {


        "degree":education,


        "university":"",

        "year":""


        }

    ]








    # -----------------
    # Projects
    # -----------------


    project_text=extract_between(

        text,

        [

        "PROJECT SUMMARY",

        "PROJECTS"

        ],

        [

        "CERTIFICATION",

        "ACHIEVEMENT"

        ]

    )



    projects=[]



    for line in project_text.split("\n"):


        if len(line)>5:


            projects.append(

                {


                "name":line,


                "description":"",


                "link":""


                }

            )



    resume["projects"]=projects







    # -----------------
    # Certifications
    # -----------------


    resume["certifications"]=[]



    # -----------------
    # Achievements
    # -----------------


    resume["achievements"]=[]





    return resume
