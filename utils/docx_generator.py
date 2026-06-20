from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tempfile




def add_heading(doc,text):

    p = doc.add_paragraph()

    run = p.add_run(text)

    run.bold=True

    run.font.size=Pt(14)

    return p





def add_text(doc,text):

    p=doc.add_paragraph()

    p.add_run(

        text

    )

    return p





def add_bullet(doc,text):

    p=doc.add_paragraph(

        style="List Bullet"

    )

    p.add_run(

        text

    )






def add_line(doc):


    p=doc.add_paragraph()

    p.add_run(

        "________________________________"

    )






def create_docx(template_name,data):


    """
    Dynamic Resume Generator

    Does not depend on placeholders
    """


    doc=Document()



    # -------------------
    # Margins
    # -------------------

    section=doc.sections[0]


    section.top_margin=Inches(0.5)

    section.bottom_margin=Inches(0.5)

    section.left_margin=Inches(0.6)

    section.right_margin=Inches(0.6)





    # -------------------
    # Header
    # -------------------


    name=data.get(

        "name",

        ""

    )


    p=doc.add_paragraph()


    p.alignment=WD_ALIGN_PARAGRAPH.CENTER


    r=p.add_run(name)


    r.bold=True

    r.font.size=Pt(20)





    title=data.get(

        "job_title",

        ""

    )


    p=doc.add_paragraph()


    p.alignment=WD_ALIGN_PARAGRAPH.CENTER


    r=p.add_run(title)


    r.font.size=Pt(12)







    contact=[]


    for key in [

        "email",

        "phone",

        "location",

        "linkedin"

    ]:


        if data.get(key):

            contact.append(

                data[key]

            )



    if contact:


        p=doc.add_paragraph()


        p.alignment=WD_ALIGN_PARAGRAPH.CENTER


        p.add_run(

            " | ".join(contact)

        )






    # -------------------
    # Summary
    # -------------------


    add_heading(

        doc,

        "PROFESSIONAL SUMMARY"

    )


    add_text(

        doc,

        data.get(

            "summary",

            ""

        )

    )







    # -------------------
    # Skills
    # -------------------


    add_heading(

        doc,

        "CORE SKILLS"

    )


    skills=data.get(

        "skills",

        []

    )


    if isinstance(skills,list):


        add_text(

            doc,

            ", ".join(skills)

        )


    else:


        add_text(

            doc,

            skills

        )






    # -------------------
    # Experience
    # -------------------


    add_heading(

        doc,

        "PROFESSIONAL EXPERIENCE"

    )



    experiences=data.get(

        "experience",

        []

    )



    for exp in experiences:


        p=doc.add_paragraph()



        r=p.add_run(

            exp.get(

                "role",

                ""

            )

        )


        r.bold=True



        p.add_run(

            " | "

            +

            exp.get(

                "company",

                ""

            )

        )




        if exp.get("duration"):


            add_text(

                doc,

                exp["duration"]

            )





        bullets=exp.get(

            "responsibilities",

            []

        )



        for b in bullets:


            add_bullet(

                doc,

                b

            )






    # -------------------
    # Education
    # -------------------


    add_heading(

        doc,

        "EDUCATION"

    )



    for edu in data.get(

        "education",

        []

    ):



        add_text(

            doc,

            edu.get(

                "degree",

                ""

            )

        )



        add_text(

            doc,

            edu.get(

                "university",

                ""

            )

        )



        add_text(

            doc,

            edu.get(

                "year",

                ""

            )

        )






    # -------------------
    # Projects
    # -------------------


    add_heading(

        doc,

        "PROJECTS"

    )



    for project in data.get(

        "projects",

        []

    ):



        add_text(

            doc,

            project.get(

                "name",

                ""

            )

        )



        add_text(

            doc,

            project.get(

                "description",

                ""

            )

        )







    # -------------------
    # Certifications
    # -------------------


    add_heading(

        doc,

        "CERTIFICATIONS"

    )



    for c in data.get(

        "certifications",

        []

    ):


        add_bullet(

            doc,

            c

        )






    # -------------------
    # Achievements
    # -------------------


    add_heading(

        doc,

        "ACHIEVEMENTS"

    )



    for a in data.get(

        "achievements",

        []

    ):


        add_bullet(

            doc,

            a

        )





    # Save


    output=tempfile.NamedTemporaryFile(

        delete=False,

        suffix=".docx"

    )



    doc.save(

        output.name

    )



    return output.name
