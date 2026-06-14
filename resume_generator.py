from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph
)

from reportlab.lib.styles import getSampleStyleSheet



def create_word_resume(text):


    path="AI_Resume.docx"


    doc=Document()


    doc.add_heading(
        "Professional Resume",
        level=1
    )


    doc.add_paragraph(text)


    doc.save(path)


    return path





def create_pdf_resume(text):


    path="AI_Resume.pdf"


    pdf=SimpleDocTemplate(path)


    styles=getSampleStyleSheet()


    pdf.build(
        [
            Paragraph(
                text.replace(
                    "\n",
                    "<br/>"
                ),
                styles["Normal"]
            )
        ]
    )


    return path