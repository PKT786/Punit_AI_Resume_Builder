"""
resume_builder.py
------------------
Parametrized .docx resume generator supporting 3 quality tiers, each with
5 style variants (15 templates total). All templates share the same
single-column, table-free, image-free layout so every one of them stays
ATS-parseable — only color, font, and header treatment change.

Usage:
    from resume_builder import build_resume, TEMPLATES, pick_random_template

    template = pick_random_template("premium")     # -> dict from TEMPLATES
    docx_bytes = build_resume(data, template)
"""

import io
import random
from docx import Document
from docx.shared import Pt, Twips, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PAGE_WIDTH_TWIPS = 12240          # US Letter
MARGIN_TWIPS = 720                # 0.5"
MARGIN = Inches(0.5)


# ============================================================== TEMPLATES ==
# header_mode:
#   "banner"    -> full-bleed colored background block behind name/title/contact (Premium)
#   "underline" -> colored name + a colored rule under the whole header block (Professional)
#   "plain"     -> black/gray name, thin gray rule, no color fill (Simple)
TEMPLATES = {
    "simple": [
        {
            "id": "simple-1", "name": "Classic Black", "tier": "simple",
            "header_mode": "plain",
            "font_head": "Calibri", "font_body": "Calibri",
            "navy": "1A1A1A", "accent": "444444", "darkgray": "222222", "midgray": "5A5A5A",
        },
        {
            "id": "simple-2", "name": "Arial Clean", "tier": "simple",
            "header_mode": "plain",
            "font_head": "Arial", "font_body": "Arial",
            "navy": "1F3864", "accent": "1F3864", "darkgray": "262626", "midgray": "595959",
        },
        {
            "id": "simple-3", "name": "Times Traditional", "tier": "simple",
            "header_mode": "underline",
            "font_head": "Times New Roman", "font_body": "Times New Roman",
            "navy": "5B1A1A", "accent": "5B1A1A", "darkgray": "222222", "midgray": "595959",
        },
        {
            "id": "simple-4", "name": "Verdana Minimal", "tier": "simple",
            "header_mode": "plain",
            "font_head": "Verdana", "font_body": "Verdana",
            "navy": "2B2B2B", "accent": "2B2B2B", "darkgray": "2B2B2B", "midgray": "5A5A5A",
        },
        {
            "id": "simple-5", "name": "Georgia Classic", "tier": "simple",
            "header_mode": "underline",
            "font_head": "Georgia", "font_body": "Calibri",
            "navy": "0E4D4D", "accent": "0E4D4D", "darkgray": "222222", "midgray": "555555",
        },
    ],
    "professional": [
        {
            "id": "professional-1", "name": "Navy Professional", "tier": "professional",
            "header_mode": "underline",
            "font_head": "Cambria", "font_body": "Calibri",
            "navy": "1F3864", "accent": "3E6B9C", "darkgray": "262626", "midgray": "595959",
        },
        {
            "id": "professional-2", "name": "Teal Corporate", "tier": "professional",
            "header_mode": "underline",
            "font_head": "Cambria", "font_body": "Calibri",
            "navy": "0F4C4C", "accent": "2E8B8B", "darkgray": "232323", "midgray": "555555",
        },
        {
            "id": "professional-3", "name": "Maroon Executive", "tier": "professional",
            "header_mode": "underline",
            "font_head": "Georgia", "font_body": "Calibri",
            "navy": "601C1C", "accent": "A15C5C", "darkgray": "262626", "midgray": "595959",
        },
        {
            "id": "professional-4", "name": "Forest Business", "tier": "professional",
            "header_mode": "underline",
            "font_head": "Cambria", "font_body": "Calibri",
            "navy": "1E4620", "accent": "8A7B4E", "darkgray": "232323", "midgray": "555555",
        },
        {
            "id": "professional-5", "name": "Slate Modern", "tier": "professional",
            "header_mode": "underline",
            "font_head": "Calibri", "font_body": "Calibri",
            "navy": "334155", "accent": "64748B", "darkgray": "262626", "midgray": "595959",
        },
    ],
    "premium": [
        {
            "id": "premium-1", "name": "Navy & Gold", "tier": "premium",
            "header_mode": "banner",
            "font_head": "Cambria", "font_body": "Calibri",
            "navy": "16294D", "accent": "9C7A2E", "darkgray": "2B2B2B", "midgray": "5A5A5A",
        },
        {
            "id": "premium-2", "name": "Charcoal & Rose Gold", "tier": "premium",
            "header_mode": "banner",
            "font_head": "Cambria", "font_body": "Calibri",
            "navy": "2B2B2B", "accent": "B76E79", "darkgray": "262626", "midgray": "595959",
        },
        {
            "id": "premium-3", "name": "Emerald & Cream", "tier": "premium",
            "header_mode": "banner",
            "font_head": "Georgia", "font_body": "Calibri",
            "navy": "0B3D2E", "accent": "C9A877", "darkgray": "232323", "midgray": "555555",
        },
        {
            "id": "premium-4", "name": "Burgundy & Gold", "tier": "premium",
            "header_mode": "banner",
            "font_head": "Cambria", "font_body": "Calibri",
            "navy": "4A1023", "accent": "C9A227", "darkgray": "262626", "midgray": "595959",
        },
        {
            "id": "premium-5", "name": "Midnight & Silver", "tier": "premium",
            "header_mode": "banner",
            "font_head": "Cambria", "font_body": "Calibri",
            "navy": "0D1B2A", "accent": "9AA5B1", "darkgray": "232323", "midgray": "555555",
        },
        {
            "id": "premium-6", "name": "Espresso & Copper", "tier": "premium",
            "header_mode": "banner",
            "font_head": "Georgia", "font_body": "Calibri",
            "navy": "2E1D14", "accent": "C77B4E", "darkgray": "262220", "midgray": "5C534D",
        },
        {
            "id": "premium-7", "name": "Royal Purple & Gold", "tier": "premium",
            "header_mode": "banner",
            "font_head": "Cambria", "font_body": "Calibri",
            "navy": "2E1A47", "accent": "D4AF37", "darkgray": "241F2B", "midgray": "5A5566",
        },
        {
            "id": "premium-8", "name": "Onyx & Platinum", "tier": "premium",
            "header_mode": "banner",
            "font_head": "Cambria", "font_body": "Calibri",
            "navy": "111214", "accent": "C7C9CC", "darkgray": "222222", "midgray": "58595B",
        },
        {
            "id": "premium-9", "name": "Deep Teal & Champagne", "tier": "premium",
            "header_mode": "banner",
            "font_head": "Georgia", "font_body": "Calibri",
            "navy": "0A3B3F", "accent": "D9C79B", "darkgray": "20272A", "midgray": "526066",
        },
    ],
}


def pick_random_template(tier: str) -> dict:
    tier = (tier or "premium").lower()
    if tier not in TEMPLATES:
        tier = "premium"
    return random.choice(TEMPLATES[tier])


def get_template_by_id(template_id: str) -> dict:
    for tier_list in TEMPLATES.values():
        for t in tier_list:
            if t["id"] == template_id:
                return t
    return TEMPLATES["premium"][0]


# --------------------------------------------------------------- oxml helpers
def _shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def _border_paragraph(paragraph, **sides):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side, spec in sides.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(spec.get('sz', 8)))
        el.set(qn('w:space'), str(spec.get('space', 4)))
        el.set(qn('w:color'), spec.get('color', '000000'))
        pBdr.append(el)
    pPr.append(pBdr)


def _char_spacing(run, points):
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), str(int(points * 20)))
    rPr.append(spacing)


def _add_tab_stop(paragraph, position_twips, alignment=WD_TAB_ALIGNMENT.RIGHT):
    paragraph.paragraph_format.tab_stops.add_tab_stop(Twips(position_twips), alignment)


# ============================================================== =========== #
class ResumeBuilder:
    """Builds a single-column, ATS-safe .docx resume using a style/template dict."""

    def __init__(self, style: dict):
        self.style = style
        self.navy = RGBColor.from_string(style["navy"])
        self.accent = RGBColor.from_string(style["accent"])
        self.darkgray = RGBColor.from_string(style["darkgray"])
        self.midgray = RGBColor.from_string(style["midgray"])
        self.white = RGBColor(0xFF, 0xFF, 0xFF)
        self.font_head = style["font_head"]
        self.font_body = style["font_body"]
        self.header_mode = style["header_mode"]

        self.doc = Document()
        normal = self.doc.styles["Normal"]
        normal.font.name = self.font_body
        normal.font.size = Pt(10.5)

        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        if self.header_mode == "banner":
            section.top_margin = Pt(0)
            section.left_margin = Pt(0)
            section.right_margin = Pt(0)
        else:
            section.top_margin = Pt(30)
            section.left_margin = MARGIN
            section.right_margin = MARGIN
        section.bottom_margin = Pt(28)

        self._simulate_margin = self.header_mode == "banner"

    # ---- low level run/paragraph helpers ---------------------------------
    def _run(self, p, text, *, font=None, size=10.5, color=None, bold=False,
              italic=False, underline=False, spacing=None):
        r = p.add_run(text)
        r.font.name = font or self.font_body
        r.font.size = Pt(size)
        r.font.color.rgb = color if color else self.darkgray
        r.font.bold = bold
        r.font.italic = italic
        r.font.underline = underline
        if spacing:
            _char_spacing(r, spacing)
        return r

    def _p(self, *, align=None, before=0, after=6, extra_left=0):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        if self._simulate_margin:
            p.paragraph_format.left_indent = MARGIN + Pt(extra_left)
            p.paragraph_format.right_indent = MARGIN
        elif extra_left:
            p.paragraph_format.left_indent = Pt(extra_left)
        if align:
            p.alignment = align
        return p

    # ---- header ------------------------------------------------------------
    def add_header(self, data):
        name = (data.get("name") or "").upper()
        title = (data.get("title") or "").upper()
        contact_bits = [b for b in [data.get("location"), data.get("phone"), data.get("email")] if b]

        if self.header_mode == "banner":
            p_name = self.doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_name.paragraph_format.space_before = Pt(26)
            p_name.paragraph_format.space_after = Pt(4)
            _shade_paragraph(p_name, self.style["navy"])
            self._run(p_name, name, font=self.font_head, size=26, color=self.white, bold=True, spacing=1.5)

            p_title = self.doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.space_after = Pt(4)
            _shade_paragraph(p_title, self.style["navy"])
            if title:
                self._run(p_title, title, font=self.font_body, size=10.5, color=self.accent, bold=True, spacing=0.6)

            p_contact = self.doc.add_paragraph()
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_contact.paragraph_format.space_before = Pt(2)
            p_contact.paragraph_format.space_after = Pt(26)
            _shade_paragraph(p_contact, self.style["navy"])
            pale = RGBColor(0xD8, 0xDC, 0xE6)
            self._run(p_contact, "   |   ".join(contact_bits), size=9, color=pale)
            if data.get("linkedin"):
                self._run(p_contact, "   |   ", size=9, color=pale)
                self._run(p_contact, "LinkedIn Profile", size=9, color=self.white, bold=True, underline=True)

        elif self.header_mode == "underline":
            p_name = self._p(align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
            self._run(p_name, name, font=self.font_head, size=22, color=self.navy, bold=True, spacing=1.0)
            if title:
                p_title = self._p(align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
                self._run(p_title, title, font=self.font_body, size=11, color=self.midgray, bold=True)
            p_contact = self._p(align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
            _border_paragraph(p_contact, bottom={"color": self.style["accent"], "sz": 12, "space": 6})
            self._run(p_contact, "   |   ".join(contact_bits), size=9.5, color=self.midgray)
            if data.get("linkedin"):
                self._run(p_contact, "   |   ", size=9.5, color=self.midgray)
                self._run(p_contact, "LinkedIn Profile", size=9.5, color=self.navy, bold=True, underline=True)

        else:  # plain
            p_name = self._p(align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
            self._run(p_name, name, font=self.font_head, size=20, color=self.navy, bold=True)
            if title:
                p_title = self._p(align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
                self._run(p_title, title, size=10.5, color=self.midgray, italic=True)
            p_contact = self._p(align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
            _border_paragraph(p_contact, bottom={"color": "BFBFBF", "sz": 6, "space": 6})
            self._run(p_contact, "   |   ".join(contact_bits), size=9.5, color=self.midgray)
            if data.get("linkedin"):
                self._run(p_contact, "   |   ", size=9.5, color=self.midgray)
                self._run(p_contact, "LinkedIn Profile", size=9.5, color=self.navy, underline=True)

    # ---- sections ------------------------------------------------------------
    def section_heading(self, text):
        p = self._p(before=16, after=6)
        _border_paragraph(p, bottom={"color": self.style["accent"], "sz": 10, "space": 4})
        self._run(p, text.upper(), font=self.font_head, size=12.5, color=self.navy, bold=True, spacing=1.1)

    def paragraph_text(self, text, justify=True):
        p = self._p(align=WD_ALIGN_PARAGRAPH.JUSTIFY if justify else None, before=4, after=8)
        self._run(p, text, size=10.5, color=self.darkgray)

    def skill_line(self, label, value):
        p = self._p(after=7)
        base = MARGIN_TWIPS if self._simulate_margin else 0
        _add_tab_stop(p, base + int(1.9 * 1440), WD_TAB_ALIGNMENT.LEFT)
        self._run(p, label, size=10.5, color=self.navy, bold=True)
        self._run(p, "\t")
        self._run(p, value, size=10.5, color=self.darkgray)

    def job_header(self, role, company, dates):
        p = self._p(before=12, after=1, extra_left=14 if self.style["tier"] != "simple" else 0)
        if self.style["tier"] != "simple":
            _border_paragraph(p, left={"color": self.style["accent"], "sz": 18, "space": 8})
        right_pos = PAGE_WIDTH_TWIPS - MARGIN_TWIPS if self._simulate_margin else PAGE_WIDTH_TWIPS - 2 * MARGIN_TWIPS
        _add_tab_stop(p, right_pos, WD_TAB_ALIGNMENT.RIGHT)
        self._run(p, role, font=self.font_head, size=12, color=self.navy, bold=True)
        if company:
            self._run(p, f"   —   {company}", font=self.font_head, size=12, color=self.darkgray, bold=True)
        self._run(p, "\t")
        self._run(p, dates, size=9.5, color=self.midgray, italic=True)

    def project_line(self, text):
        p = self._p(after=5, extra_left=14 if self.style["tier"] != "simple" else 0)
        self._run(p, "Project: ", size=10, color=self.accent, bold=True, italic=True)
        self._run(p, text, size=10, color=self.midgray, italic=True)

    def bullet(self, text):
        p = self.doc.add_paragraph()
        base = MARGIN if self._simulate_margin else Pt(0)
        p.paragraph_format.left_indent = base + Pt(24)
        p.paragraph_format.right_indent = MARGIN if self._simulate_margin else Pt(0)
        p.paragraph_format.first_line_indent = Pt(-12)
        p.paragraph_format.space_after = Pt(5)
        self._run(p, "•  ", size=10.5, color=self.accent, bold=True)
        self._run(p, text, size=10.5, color=self.darkgray)

    def edu_entry(self, degree, details_line):
        if degree:
            p = self._p(after=1)
            self._run(p, degree, size=10.5, color=self.navy, bold=True)
        if details_line:
            p2 = self._p(after=10)
            self._run(p2, details_line, size=10.5, color=self.darkgray)

    def to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf.read()


# --------------------------------------------------------------------- main
def build_resume(data: dict, template: dict = None, tier: str = "premium") -> bytes:
    """
    data: resume content dict (see resume_parser output / app.py form)
    template: a specific template dict from TEMPLATES (overrides `tier`)
    tier: "simple" | "professional" | "premium" — used to pick a random
          template from that tier if `template` isn't supplied
    """
    style = template or pick_random_template(tier)
    rb = ResumeBuilder(style)
    rb.add_header(data)

    if data.get("summary"):
        rb.section_heading("Professional Summary")
        rb.paragraph_text(data["summary"])

    if data.get("skills"):
        rb.section_heading("Technical Skills")
        for s in data["skills"]:
            if s.get("label") and s.get("value"):
                rb.skill_line(s["label"], s["value"])

    if data.get("experience"):
        rb.section_heading("Professional Experience")
        for job in data["experience"]:
            rb.job_header(job.get("role", ""), job.get("company", ""), job.get("dates", ""))
            if job.get("project"):
                rb.project_line(job["project"])
            for b in job.get("bullets", []):
                if b.strip():
                    rb.bullet(b.strip())

    if data.get("certifications"):
        rb.section_heading("Certifications")
        for c in data["certifications"]:
            if c.strip():
                rb.bullet(c.strip())

    if data.get("awards"):
        rb.section_heading("Awards & Achievements")
        for a in data["awards"]:
            if a.strip():
                rb.bullet(a.strip())

    if data.get("education"):
        rb.section_heading("Education")
        for edu in data["education"]:
            details_line = " | ".join([x for x in [edu.get("school"), edu.get("details")] if x])
            rb.edu_entry(edu.get("degree", ""), details_line)

    return rb.to_bytes()


# ------------------------------------------------------------------ sample
SAMPLE_DATA = {
    "name": "Jane Doe",
    "title": "Senior Data Engineer | Cloud & Analytics Specialist",
    "location": "Austin, TX",
    "phone": "+1-555-123-4567",
    "email": "jane.doe@email.com",
    "linkedin": "https://www.linkedin.com/in/janedoe/",
    "summary": "Data Engineer with 8+ years of experience building scalable pipelines and cloud-native analytics platforms across fintech and retail domains.",
    "skills": [
        {"label": "Languages", "value": "Python, SQL, Scala"},
        {"label": "Cloud", "value": "AWS, GCP, Azure"},
    ],
    "experience": [
        {
            "role": "Senior Data Engineer", "company": "Acme Corp", "dates": "Jan 2022 – Present",
            "project": "Retail Analytics Platform",
            "bullets": [
                "Built streaming pipelines processing 2TB/day using Spark and Kafka.",
                "Reduced query costs by 40% through warehouse optimization.",
            ],
        }
    ],
    "certifications": ["AWS Certified Solutions Architect"],
    "awards": ["Employee of the Year — 2023"],
    "education": [
        {"degree": "B.S. Computer Science", "school": "University of Texas", "details": "2015"}
    ],
}

if __name__ == "__main__":
    for tier in ["simple", "professional", "premium"]:
        t = TEMPLATES[tier][0]
        out = build_resume(SAMPLE_DATA, template=t)
        fname = f"sample_{t['id']}.docx"
        with open(fname, "wb") as f:
            f.write(out)
        print("wrote", fname)
